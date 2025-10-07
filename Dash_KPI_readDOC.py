import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import docx
import io
import re
from collections import defaultdict

# --- Configuração da Página ---
st.set_page_config(
    page_title="Dashboard de Qualidade Dinâmico",
    page_icon="📊",
    layout="wide"
)

# --- Estilo CSS Customizado ---
st.markdown("""
<style>
    .stMetric {
        background-color: #1F2937;
        border: 1px solid #374151;
        border-radius: 0.75rem;
        padding: 1.5rem;
    }
    .stMetric:hover {
        border-color: #556070;
    }
    .stPlotlyChart {
        background-color: #1F2937;
        border: 1px solid #374151;
        border-radius: 0.75rem;
        padding: 1.5rem;
    }
</style>
""", unsafe_allow_html=True)

# --- Cores para os Gráficos ---
CHART_COLORS = {
    'Passou': 'rgba(74, 222, 128, 0.8)',
    'Falhou': 'rgba(248, 113, 113, 0.8)',
    'Bloqueado': 'rgba(251, 146, 60, 0.8)',
    'Não Executado': 'rgba(156, 163, 175, 0.6)'
}

# --- Funções de Extração de Dados do .DOC ---

def parse_status(status_text):
    """Normaliza o texto do status."""
    status_text = status_text.lower()
    if 'passou' in status_text:
        return 'Passou'
    if 'falhado' in status_text or 'falhou' in status_text:
        return 'Falhou'
    if 'bloqueado' in status_text:
        return 'Bloqueado'
    if 'não executado' in status_text:
        return 'Não Executado'
    return None

def extract_data_from_doc(doc_file):
    """
    Extrai e processa dados de teste de um arquivo .doc do TestLink.
    """
    try:
        document = docx.Document(doc_file)
        
        test_data = {
            'web': defaultdict(int),
            'android': defaultdict(int),
            'ios': defaultdict(int)
        }
        bug_impact_data = defaultdict(int)

        current_platform = None

        # A lógica assume que as tabelas seguem os cabeçalhos da plataforma
        doc_tables = iter(document.tables)
        for para in document.paragraphs:
            text = para.text.lower()
            if 'plataforma: mobile android' in text:
                current_platform = 'android'
            elif 'plataforma: mobile ios' in text:
                current_platform = 'ios'
            elif 'plataforma: web' in text:
                current_platform = 'web'
            
            # Heurística para verificar se um parágrafo é seguido por uma tabela de caso de teste
            if 'caso de teste' in text and current_platform:
                try:
                    table = next(doc_tables)
                    status_text = ''
                    comment_text = ''
                    for row in table.rows:
                        if 'Resultado da Execução' in row.cells[0].text and len(row.cells) > 1:
                            status_text = row.cells[-1].text
                        if 'Comentários' in row.cells[0].text and len(row.cells) > 1:
                            comment_text = row.cells[-1].text
                    
                    status = parse_status(status_text)
                    if status:
                        test_data[current_platform][status] += 1
                        if status in ['Falhou', 'Bloqueado'] and comment_text:
                            bug_match = re.search(r'(PH-\d+.*?)(?=\s\s|$)', comment_text)
                            if bug_match:
                                bug_id = bug_match.group(1).strip()
                                bug_impact_data[bug_id] += 1
                except (StopIteration, IndexError):
                    continue

        # Pós-processamento para testes não executados com base na estrutura do documento
        total_tcs_per_platform = 20
        test_data['android']['Não Executado'] = total_tcs_per_platform - sum(test_data['android'].values())
        test_data['ios']['Não Executado'] = total_tcs_per_platform - sum(test_data['ios'].values())
        test_data['web']['Não Executado'] = total_tcs_per_platform - sum(test_data['web'].values())
        
        # Converte defaultdicts para dicts normais
        final_test_data = {k: dict(v) for k, v in test_data.items()}
        final_bug_data = dict(bug_impact_data)
        
        return final_test_data, final_bug_data
    except Exception as e:
        st.error(f"Erro ao processar o arquivo .doc: {e}")
        return None, None

# --- Função Principal da UI ---
def run_dashboard(test_data, bug_impact_data):
    """Renderiza o dashboard completo com base nos dados fornecidos."""
    # Cabeçalho
    st.title("Dashboard de Qualidade")
    st.markdown("Projeto Payment Hub")
    st.markdown("---")

    # --- Cálculos de KPIs ---
    total_testes = sum(sum(platform.values()) for platform in test_data.values())
    
    executados = 0
    for platform_data in test_data.values():
        for status, count in platform_data.items():
            if status != 'Não Executado':
                executados += count
    
    total_passou = sum(d.get('Passou', 0) for d in test_data.values())
    total_falhou = sum(d.get('Falhou', 0) for d in test_data.values())
    total_bloqueado = sum(d.get('Bloqueado', 0) for d in test_data.values())

    cobertura = (executados / total_testes) * 100 if total_testes > 0 else 0
    taxa_sucesso = (total_passou / executados) * 100 if executados > 0 else 0
    taxa_defeito = (total_falhou / executados) * 100 if executados > 0 else 0
    taxa_bloqueio = (total_bloqueado / executados) * 100 if executados > 0 else 0
    
    # Seção de KPIs
    kpi_cols = st.columns(6)
    kpi_cols[0].metric("Total de Testes", f"{total_testes}")
    kpi_cols[1].metric("Testes Executados", f"{executados}")
    kpi_cols[2].metric("Cobertura de Teste", f"{cobertura:.1f}%")
    kpi_cols[3].metric("Taxa de Sucesso", f"{taxa_sucesso:.1f}%", help="Percentual de testes que passaram em relação aos executados.")
    kpi_cols[4].metric("Taxa de Defeito", f"{taxa_defeito:.1f}%", help="Percentual de testes que falharam em relação aos executados.")
    kpi_cols[5].metric("Taxa de Bloqueio", f"{taxa_bloqueio:.1f}%", help="Percentual de testes bloqueados em relação aos executados.")

    st.markdown("<br>", unsafe_allow_html=True)

    # Seção de Gráficos
    # Linha 1: Status Geral e Status por Plataforma
    chart_cols_top = st.columns([2, 3])

    with chart_cols_top[0]:
        st.subheader("Status Geral de Execução")
        total_not_executed = sum(d.get('Não Executado', 0) for d in test_data.values())
        
        fig_geral = go.Figure(data=[go.Pie(
            labels=['Passou', 'Falhou', 'Bloqueado', 'Não Executado'],
            values=[total_passou, total_falhou, total_bloqueado, total_not_executed],
            hole=.4,
            marker_colors=[CHART_COLORS.get(s, '#CCCCCC') for s in ['Passou', 'Falhou', 'Bloqueado', 'Não Executado']],
            textinfo='percent+label'
        )])
        fig_geral.update_layout(showlegend=False, margin=dict(t=0, b=0, l=0, r=0), paper_bgcolor='#1F2937', font_color='white')
        st.plotly_chart(fig_geral, use_container_width=True)

    with chart_cols_top[1]:
        st.subheader("Status por Plataforma")
        platforms = ['WEB', 'Mobile Android', 'MOBILE iOS']
        statuses = ['Passou', 'Falhou', 'Bloqueado', 'Não Executado']
        
        fig_plataforma = go.Figure()
        for status in statuses:
            values = [test_data.get(p_key, {}).get(status, 0) for p_key in ['web', 'android', 'ios']]
            fig_plataforma.add_trace(go.Bar(name=status, x=platforms, y=values, marker_color=CHART_COLORS.get(status, '#CCCCCC')))
        
        fig_plataforma.update_layout(barmode='stack', legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1), xaxis=dict(tickfont=dict(color='white')), yaxis=dict(tickfont=dict(color='white')), paper_bgcolor='#1F2937', plot_bgcolor='#1F2937', font_color='white', margin=dict(t=20, b=20, l=20, r=20))
        st.plotly_chart(fig_plataforma, use_container_width=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Linha 2: Impacto de Bugs
    if bug_impact_data:
        st.subheader("Impacto de Bugs nos Casos de Teste")
        bug_labels = list(bug_impact_data.keys())
        bug_values = list(bug_impact_data.values())

        fig_bugs = go.Figure(go.Bar(
            x=bug_values, y=bug_labels, orientation='h',
            marker=dict(color='rgba(239, 68, 68, 0.7)', line=dict(color='rgb(239, 68, 68)', width=1))
        ))
        fig_bugs.update_layout(xaxis_title="Casos de Teste Impactados", yaxis=dict(autorange="reversed"), paper_bgcolor='#1F2937', plot_bgcolor='#1F2937', font_color='white', margin=dict(t=20, b=20, l=20, r=20))
        st.plotly_chart(fig_bugs, use_container_width=True)


# --- Aplicação Principal ---
def main():
    st.sidebar.header("📁 Carregar Relatório")
    st.sidebar.info("Para que o script funcione, instale a biblioteca `python-docx`.")
    uploaded_file = st.sidebar.file_uploader(
        "Selecione o arquivo de relatório (.doc ou .docx)",
        type=['doc', 'docx']
    )
    
    if uploaded_file is not None:
        # Quando um arquivo é carregado, processa-o
        file_buffer = io.BytesIO(uploaded_file.getvalue())
        test_data, bug_impact_data = extract_data_from_doc(file_buffer)
        
        if test_data and bug_impact_data is not None:
            run_dashboard(test_data, bug_impact_data)
        else:
            st.error("Não foi possível extrair dados do arquivo. Verifique o formato.")
    else:
        # Se nenhum arquivo for carregado, mostra os dados de exemplo
        st.info("👈 Por favor, carregue um arquivo de relatório .doc para começar.")
        
        # Dados de exemplo (os dados originais do seu relatório)
        example_test_data = {
            'web': {'Passou': 3, 'Falhou': 3, 'Bloqueado': 8, 'Não Executado': 6},
            'android': {'Passou': 0, 'Falhou': 0, 'Bloqueado': 0, 'Não Executado': 20},
            'ios': {'Passou': 0, 'Falhou': 0, 'Bloqueado': 0, 'Não Executado': 20}
        }
        example_bug_data = {
            'PH-177 [QA] Erro 500 ao finalizar pedido em QAS': 8,
            'PH-178 (QA)Pedido não exibe forma de pagamento...': 1,
            'PH-179 (QA) Erro na criação do pedido...': 1
        }
        run_dashboard(example_test_data, example_bug_data)

if __name__ == "__main__":
    main()
